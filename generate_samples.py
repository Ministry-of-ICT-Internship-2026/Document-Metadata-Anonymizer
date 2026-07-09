"""
generate_samples.py — builds small sample .docx files with realistic FAKE
metadata (fake names, fake company, fake comments) for exercising
docx_stripper.py in tests/test_docx.py.

Run:
    python tests/sample_files/generate_samples.py

Produces (next to this script):
    plain_metadata.docx   - populated core_properties, no comments/tracking
    with_comments.docx    - populated core_properties + Word comments
    with_tracked_changes.docx - populated core_properties + w:ins/w:del
    no_metadata.docx      - core_properties left at python-docx defaults
"""

from __future__ import annotations

import copy
import os
import zipfile
from datetime import datetime

from docx import Document
from lxml import etree

HERE = os.path.dirname(os.path.abspath(__file__))

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

NSMAP = {"w": W_NS}


def _q(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


EXTENDED_PROPS_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"


def _set_fake_core_properties(document: Document) -> None:
    props = document.core_properties
    props.author = "John Doe"
    props.last_modified_by = "jdoe_mdaict"
    props.title = "Draft Circular on Data Protection"
    props.subject = "Internal Policy"
    props.keywords = "confidential, draft, mdaict"
    props.comments = "Reviewed by legal, pending sign-off"
    props.category = "Policy"
    props.revision = 4
    props.created = datetime(2024, 3, 1, 10, 0, 0)
    props.modified = datetime(2024, 6, 12, 14, 22, 0)


def _set_fake_company(docx_path: str) -> None:
    """Company lives in docProps/app.xml, not core.xml, so python-docx's
    core_properties can't set it — write it directly."""
    tmp_path = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        app_xml = zin.read("docProps/app.xml")
        other_parts = {n: zin.read(n) for n in names if n != "docProps/app.xml"}

    root = etree.fromstring(app_xml)
    company_el = root.find(f"{{{EXTENDED_PROPS_NS}}}Company")
    if company_el is None:
        company_el = etree.SubElement(root, f"{{{EXTENDED_PROPS_NS}}}Company")
    company_el.text = "Ministry of ICT"
    new_app_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other_parts.items():
            zout.writestr(name, data)
        zout.writestr("docProps/app.xml", new_app_xml)
    os.replace(tmp_path, docx_path)


def _make_base_document(body_paragraphs) -> Document:
    document = Document()
    for text in body_paragraphs:
        document.add_paragraph(text)
    _set_fake_core_properties(document)
    return document


def make_plain_metadata_sample() -> str:
    path = os.path.join(HERE, "plain_metadata.docx")
    document = _make_base_document(
        [
            "Ministry of ICT and National Guidance",
            "Subject: Draft Circular on Data Protection Practices",
            "This circular outlines minimum data-handling standards for MDA staff.",
        ]
    )
    document.save(path)
    _set_fake_company(path)
    return path


def make_no_metadata_sample() -> str:
    path = os.path.join(HERE, "no_metadata.docx")
    document = Document()
    document.add_paragraph("A short memo with no populated metadata fields.")
    document.save(path)
    return path


def _inject_comment(docx_path: str) -> None:
    """Add a single Word comment authored by a fake reviewer, anchored to
    the first run of the first paragraph. Requires touching six parts, so
    this is done by hand-editing the already-saved .docx's zip."""
    tmp_path = docx_path + ".tmp"

    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        document_xml = zin.read("word/document.xml")
        content_types_xml = zin.read("[Content_Types].xml")
        rels_xml = zin.read("word/_rels/document.xml.rels")
        other_parts = {
            n: zin.read(n) for n in names
            if n not in ("word/document.xml", "[Content_Types].xml", "word/_rels/document.xml.rels")
        }

    # --- comments.xml ---
    comments_root = etree.Element(_q("comments"), nsmap=NSMAP)
    comment = etree.SubElement(comments_root, _q("comment"))
    comment.set(_q("id"), "0")
    comment.set(_q("author"), "Jane K.")
    comment.set(_q("initials"), "JK")
    comment.set(_q("date"), "2024-06-10T09:15:00Z")
    p = etree.SubElement(comment, _q("p"))
    r = etree.SubElement(p, _q("r"))
    t = etree.SubElement(r, _q("t"))
    t.text = "Please confirm this figure with Finance."
    comments_bytes = etree.tostring(comments_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    # --- document.xml: wrap first run of first paragraph with comment range ---
    doc_root = etree.fromstring(document_xml)
    body = doc_root.find(_q("body"))
    first_paragraph = body.find(_q("p"))
    first_run = first_paragraph.find(_q("r"))

    range_start = etree.Element(_q("commentRangeStart"))
    range_start.set(_q("id"), "0")
    range_end = etree.Element(_q("commentRangeEnd"))
    range_end.set(_q("id"), "0")
    ref_run = etree.Element(_q("r"))
    ref = etree.SubElement(ref_run, _q("commentReference"))
    ref.set(_q("id"), "0")

    first_paragraph.insert(list(first_paragraph).index(first_run), range_start)
    first_paragraph.insert(list(first_paragraph).index(first_run) + 1, range_end)
    first_paragraph.insert(list(first_paragraph).index(range_end) + 1, ref_run)
    new_document_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    # --- [Content_Types].xml: register comments part ---
    ct_root = etree.fromstring(content_types_xml)
    override = etree.SubElement(ct_root, f"{{{CT_NS}}}Override")
    override.set("PartName", "/word/comments.xml")
    override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
    new_ct_xml = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    # --- word/_rels/document.xml.rels: relate document to comments ---
    rels_root = etree.fromstring(rels_xml)
    rel = etree.SubElement(rels_root, f"{{{RELS_NS}}}Relationship")
    rel.set("Id", "rIdComments1")
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments")
    rel.set("Target", "comments.xml")
    new_rels_xml = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other_parts.items():
            zout.writestr(name, data)
        zout.writestr("word/document.xml", new_document_xml)
        zout.writestr("[Content_Types].xml", new_ct_xml)
        zout.writestr("word/_rels/document.xml.rels", new_rels_xml)
        zout.writestr("word/comments.xml", comments_bytes)

    os.replace(tmp_path, docx_path)


def make_with_comments_sample() -> str:
    path = os.path.join(HERE, "with_comments.docx")
    document = _make_base_document(
        [
            "Ministry of ICT and National Guidance",
            "Budget line 4.2 allocates UGX 120,000,000 to the pilot rollout.",
        ]
    )
    document.save(path)
    _set_fake_company(path)
    _inject_comment(path)
    return path


def _inject_tracked_change(docx_path: str) -> None:
    """Wrap the first paragraph's run in <w:ins> attributed to a fake
    reviewer, simulating an accepted-looking edit that still carries
    tracked-change metadata in the XML."""
    tmp_path = docx_path + ".tmp"

    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        document_xml = zin.read("word/document.xml")
        other_parts = {n: zin.read(n) for n in names if n != "word/document.xml"}

    doc_root = etree.fromstring(document_xml)
    body = doc_root.find(_q("body"))
    first_paragraph = body.find(_q("p"))
    first_run = first_paragraph.find(_q("r"))
    run_index = list(first_paragraph).index(first_run)
    first_paragraph.remove(first_run)

    ins = etree.Element(_q("ins"))
    ins.set(_q("id"), "1")
    ins.set(_q("author"), "John Doe")
    ins.set(_q("date"), "2024-06-11T08:00:00Z")
    ins.append(first_run)
    first_paragraph.insert(run_index, ins)

    new_document_xml = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other_parts.items():
            zout.writestr(name, data)
        zout.writestr("word/document.xml", new_document_xml)

    os.replace(tmp_path, docx_path)


def make_with_tracked_changes_sample() -> str:
    path = os.path.join(HERE, "with_tracked_changes.docx")
    document = _make_base_document(
        [
            "Ministry of ICT and National Guidance",
            "All MDA staff must complete the data-protection training by Q3.",
        ]
    )
    document.save(path)
    _set_fake_company(path)
    _inject_tracked_change(path)
    return path


def _inject_attached_template(docx_path: str) -> None:
    """Add a <w:attachedTemplate> reference in word/settings.xml pointing at
    a fake internal server path, simulating a letterhead .dotx that leaks
    an internal file-share path."""
    tmp_path = docx_path + ".tmp"

    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        settings_xml = zin.read("word/settings.xml")
        rels_name = "word/_rels/settings.xml.rels"
        rels_xml = zin.read(rels_name) if rels_name in names else None
        other_parts = {
            n: zin.read(n) for n in names if n not in ("word/settings.xml", rels_name)
        }

    settings_root = etree.fromstring(settings_xml)
    att = etree.Element(_q("attachedTemplate"))
    att.set(f"{{{R_NS}}}id", "rIdTemplate1")
    settings_root.insert(0, att)
    new_settings_xml = etree.tostring(
        settings_root, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    if rels_xml is not None:
        rels_root = etree.fromstring(rels_xml)
    else:
        rels_root = etree.Element(f"{{{RELS_NS}}}Relationships")
    rel = etree.SubElement(rels_root, f"{{{RELS_NS}}}Relationship")
    rel.set("Id", "rIdTemplate1")
    rel.set(
        "Type",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/attachedTemplate",
    )
    rel.set("Target", r"file:///\\SERVER01\templates\mda_letterhead.dotx")
    rel.set("TargetMode", "External")
    new_rels_xml = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other_parts.items():
            zout.writestr(name, data)
        zout.writestr("word/settings.xml", new_settings_xml)
        zout.writestr(rels_name, new_rels_xml)

    os.replace(tmp_path, docx_path)


def make_with_attached_template_sample() -> str:
    path = os.path.join(HERE, "with_attached_template.docx")
    document = _make_base_document(
        [
            "Ministry of ICT and National Guidance",
            "This letter was produced from the standard MDA letterhead template.",
        ]
    )
    document.save(path)
    _set_fake_company(path)
    _inject_attached_template(path)
    return path


if __name__ == "__main__":
    created = [
        make_plain_metadata_sample(),
        make_no_metadata_sample(),
        make_with_comments_sample(),
        make_with_tracked_changes_sample(),
        make_with_attached_template_sample(),
    ]
    for path in created:
        print("wrote", path)
