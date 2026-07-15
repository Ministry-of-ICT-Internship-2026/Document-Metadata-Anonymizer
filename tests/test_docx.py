"""
test_docx.py — pytest suite for core/docx_stripper.py.

Run from the project root:
    pip install pytest python-docx lxml --break-system-packages
    python tests/sample_files/generate_samples.py   # (re)generate fixtures
    pytest tests/test_docx.py -v
"""

import os
import sys
import zipfile

import pytest
from docx import Document
from lxml import etree

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core.docx_stripper import (  # noqa: E402
    DocxMetadataError,
    read_metadata,
    strip_metadata,
)

SAMPLES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sample_files")


def _sample(name: str) -> str:
    path = os.path.join(SAMPLES_DIR, name)
    if not os.path.exists(path):
        pytest.skip(
            f"{name} not found — run `python tests/sample_files/generate_samples.py` first"
        )
    return path


def _paragraph_texts(path: str) -> list:
    return [p.text for p in Document(path).paragraphs]


# --------------------------------------------------------------------------
# T07 / T08 — core_properties read + clear
# --------------------------------------------------------------------------

def test_read_metadata_returns_populated_core_properties(tmp_path):
    path = _sample("plain_metadata.docx")
    meta = read_metadata(path)

    assert meta["author"] == "John Doe"
    assert meta["last_modified_by"] == "jdoe_mdaict"
    assert meta["company"] == "Ministry of ICT"
    assert meta["title"] == "Draft Circular on Data Protection"
    assert meta["subject"] == "Internal Policy"
    assert meta["category"] == "Policy"
    assert meta["revision"] == 4
    assert meta["created"] is not None
    assert meta["modified"] is not None
    # Always present, even when empty, so callers can rely on the key existing.
    for key in ("comment_authors", "tracked_change_authors", "template_path"):
        assert key in meta


def test_strip_metadata_clears_all_core_properties(tmp_path):
    src = _sample("plain_metadata.docx")
    out = str(tmp_path / "stripped.docx")

    removed = strip_metadata(src, out)

    # Report of what was removed should reflect what was actually there.
    assert removed["author"] == "John Doe"
    assert removed["company"] == "Ministry of ICT"
    assert removed["revision"] == 4

    after = read_metadata(out)
    assert after["author"] == ""
    assert after["last_modified_by"] == ""
    assert after["company"] is None
    assert after["title"] == ""
    assert after["subject"] == ""
    assert after["category"] == ""
    assert after["keywords"] == ""
    assert after["comments"] == ""
    assert after["created"] is None
    assert after["modified"] is None
    # Schema requires cp:revision to be a positive integer, so "cleared"
    # means reset to the minimum valid value, not left at the original.
    assert after["revision"] == 1


# --------------------------------------------------------------------------
# T09 — comments
# --------------------------------------------------------------------------

def test_read_metadata_finds_comment_authors():
    meta = read_metadata(_sample("with_comments.docx"))
    assert "Jane K." in meta["comment_authors"]


def test_strip_metadata_removes_comment_author_names(tmp_path):
    src = _sample("with_comments.docx")
    out = str(tmp_path / "stripped.docx")

    removed = strip_metadata(src, out)
    assert "Jane K." in removed["comment_authors"]

    after = read_metadata(out)
    assert after["comment_authors"] == []

    # The comment element itself should still exist (we scrub identity, we
    # don't need to delete the annotation), just with the author blanked.
    with zipfile.ZipFile(out) as zf:
        assert "word/comments.xml" in zf.namelist()
        root = etree.fromstring(zf.read("word/comments.xml"))
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        comments = root.findall("w:comment", ns)
        assert len(comments) == 1
        assert comments[0].get(f"{{{ns['w']}}}author") == ""


# --------------------------------------------------------------------------
# T09 — tracked changes
# --------------------------------------------------------------------------

def test_read_metadata_finds_tracked_change_authors():
    meta = read_metadata(_sample("with_tracked_changes.docx"))
    assert "John Doe" in meta["tracked_change_authors"]


def test_strip_metadata_removes_tracked_change_author_names(tmp_path):
    src = _sample("with_tracked_changes.docx")
    out = str(tmp_path / "stripped.docx")

    removed = strip_metadata(src, out)
    assert "John Doe" in removed["tracked_change_authors"]

    after = read_metadata(out)
    assert after["tracked_change_authors"] == []


# --------------------------------------------------------------------------
# T09 — attached template path
# --------------------------------------------------------------------------

def test_read_metadata_finds_attached_template_path():
    meta = read_metadata(_sample("with_attached_template.docx"))
    assert meta["template_path"] is not None
    assert "mda_letterhead.dotx" in meta["template_path"]


def test_strip_metadata_removes_attached_template_path(tmp_path):
    src = _sample("with_attached_template.docx")
    out = str(tmp_path / "stripped.docx")

    removed = strip_metadata(src, out)
    assert removed["template_path"] is not None
    assert "mda_letterhead.dotx" in removed["template_path"]

    after = read_metadata(out)
    assert after["template_path"] is None


# --------------------------------------------------------------------------
# Zero-metadata file must not crash
# --------------------------------------------------------------------------

def test_no_metadata_file_does_not_crash(tmp_path):
    src = _sample("no_metadata.docx")
    out = str(tmp_path / "stripped.docx")

    meta = read_metadata(src)
    assert meta["comment_authors"] == []
    assert meta["tracked_change_authors"] == []
    assert meta["template_path"] is None

    removed = strip_metadata(src, out)
    assert os.path.exists(out)
    assert removed["comment_authors"] == []


# --------------------------------------------------------------------------
# Visible content / formatting integrity
# --------------------------------------------------------------------------

@pytest.mark.parametrize(
    "sample_name",
    [
        "plain_metadata.docx",
        "with_comments.docx",
        "with_tracked_changes.docx",
        "with_attached_template.docx",
        "no_metadata.docx",
    ],
)
def test_visible_text_unchanged_after_stripping(tmp_path, sample_name):
    src = _sample(sample_name)
    out = str(tmp_path / f"stripped_{sample_name}")

    before_text = _paragraph_texts(src)
    strip_metadata(src, out)
    after_text = _paragraph_texts(out)

    assert before_text == after_text


@pytest.mark.parametrize(
    "sample_name",
    ["plain_metadata.docx", "with_comments.docx", "with_tracked_changes.docx"],
)
def test_paragraph_formatting_unchanged_after_stripping(tmp_path, sample_name):
    src = _sample(sample_name)
    out = str(tmp_path / f"stripped_{sample_name}")
    strip_metadata(src, out)

    before_doc = Document(src)
    after_doc = Document(out)

    assert len(before_doc.paragraphs) == len(after_doc.paragraphs)
    for before_p, after_p in zip(before_doc.paragraphs, after_doc.paragraphs):
        assert before_p.style.name == after_p.style.name
        assert len(before_p.runs) == len(after_p.runs)
        for before_r, after_r in zip(before_p.runs, after_p.runs):
            assert before_r.bold == after_r.bold
            assert before_r.italic == after_r.italic
            assert before_r.underline == after_r.underline


def test_output_is_a_valid_docx_openable_by_python_docx(tmp_path):
    src = _sample("plain_metadata.docx")
    out = str(tmp_path / "stripped.docx")
    strip_metadata(src, out)

    # Should not raise, and should still be a well-formed OPC package.
    doc = Document(out)
    assert doc.paragraphs


# --------------------------------------------------------------------------
# Error handling
# --------------------------------------------------------------------------

def test_read_metadata_raises_for_missing_file():
    with pytest.raises(DocxMetadataError):
        read_metadata("/tmp/does-not-exist-xyz.docx")


def test_read_metadata_raises_for_corrupted_file(tmp_path):
    bad_path = tmp_path / "not_a_docx.docx"
    bad_path.write_bytes(b"this is not a zip file at all")

    with pytest.raises(DocxMetadataError):
        read_metadata(str(bad_path))


def test_strip_metadata_raises_for_corrupted_file(tmp_path):
    bad_path = tmp_path / "not_a_docx.docx"
    bad_path.write_bytes(b"PK\x03\x04garbage-not-a-real-zip-entry")
    out_path = tmp_path / "out.docx"

    with pytest.raises(DocxMetadataError):
        strip_metadata(str(bad_path), str(out_path))
